from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Documentos" / "acif104_s6_MMedel.docx"
BACKUP_PATH = ROOT / "Documentos" / "acif104_s6_MMedel.backup_antes_figuras.docx"
ASSETS_DIR = ROOT / ".generated_assets"


def fmt(value) -> str:
    try:
        return f"{float(value):.4f}"
    except (TypeError, ValueError):
        return str(value)


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"No se encontró el párrafo: {text}")


def move_before_reference(reference_paragraph, element) -> None:
    reference_paragraph._p.addprevious(element)


def add_paragraph_before(document: Document, reference_paragraph, text: str = "", style: str | None = None):
    paragraph = document.add_paragraph(text)
    if style:
        paragraph.style = style
    element = paragraph._p
    element.getparent().remove(element)
    move_before_reference(reference_paragraph, element)
    return paragraph


def add_table_before(document: Document, reference_paragraph, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    element = table._tbl
    element.getparent().remove(element)
    move_before_reference(reference_paragraph, element)


def add_picture_before(document: Document, reference_paragraph, image_path: Path, caption: str, width_inches: float = 6.2) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    element = paragraph._p
    element.getparent().remove(element)
    move_before_reference(reference_paragraph, element)

    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.runs[0].italic = True
    caption_element = caption_paragraph._p
    caption_element.getparent().remove(caption_element)
    move_before_reference(reference_paragraph, caption_element)


def main() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    comparison = pd.read_csv(ASSETS_DIR / "comparacion_modelos_ml_dl.csv")
    balance = pd.read_csv(ASSETS_DIR / "comparacion_balanceo.csv")
    final_metrics = pd.read_csv(ASSETS_DIR / "metricas_modelo_final.csv").iloc[0]

    document = Document(DOCX_PATH)

    if any("Evidencia experimental generada desde el notebook" in p.text for p in document.paragraphs):
        print("La sección de evidencia ya existe en el documento. No se insertó de nuevo.")
        return

    reference = find_paragraph(document, "CONCLUSIÓN GENERAL")

    add_paragraph_before(document, reference, "Evidencia experimental generada desde el notebook", "Heading 2")
    add_paragraph_before(
        document,
        reference,
        "Los resultados siguientes fueron extraídos de los archivos generados al ejecutar el notebook completo. Esta evidencia complementa la comparación metodológica, ya que permite respaldar la selección final del modelo con métricas, gráficos y artefactos reproducibles.",
    )

    add_paragraph_before(document, reference, "Resumen del modelo final", "Heading 3")
    add_table_before(
        document,
        reference,
        ["Modelo", "Threshold", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            [
                final_metrics["Modelo"],
                fmt(final_metrics["Threshold"]),
                fmt(final_metrics["Accuracy"]),
                fmt(final_metrics["Precision"]),
                fmt(final_metrics["Recall"]),
                fmt(final_metrics["F1"]),
                fmt(final_metrics["ROC_AUC"]),
            ]
        ],
    )

    add_paragraph_before(document, reference, "Comparación ampliada de modelos", "Heading 3")
    model_rows = [
        [
            row["Modelo"],
            fmt(row["Accuracy"]),
            fmt(row["Precision"]),
            fmt(row["Recall"]),
            fmt(row["F1"]),
            fmt(row["ROC_AUC"]),
            fmt(row["Tiempo_entrenamiento_seg"]),
        ]
        for _, row in comparison.iterrows()
    ]
    add_table_before(
        document,
        reference,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "Tiempo (s)"],
        model_rows,
    )
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "grafico_modelos.png",
        "Figura 11: Comparación de modelos ML y MLP según F1-score.",
    )

    add_paragraph_before(
        document,
        reference,
        "La comparación muestra que Random Forest mantiene el mejor F1-score entre los modelos evaluados. Gradient Boosting se mantiene competitivo, mientras que las tres configuraciones MLP alcanzan resultados razonables, pero no superan el equilibrio obtenido por Random Forest en datos tabulares.",
    )

    add_paragraph_before(document, reference, "Balanceo de clases y ajuste de threshold", "Heading 3")
    balance_rows = [
        [
            row["Tecnica_balanceo"],
            fmt(row["Precision"]),
            fmt(row["Recall"]),
            fmt(row["F1"]),
            fmt(row["ROC_AUC"]),
        ]
        for _, row in balance.iterrows()
    ]
    add_table_before(
        document,
        reference,
        ["Técnica", "Precision", "Recall", "F1", "ROC-AUC"],
        balance_rows,
    )
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "grafico_balanceo.png",
        "Figura 12: Impacto de las técnicas de balanceo sobre Random Forest.",
    )
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "threshold_plot.png",
        "Figura 13: Evaluación de thresholds para precision, recall y F1-score.",
    )
    add_paragraph_before(
        document,
        reference,
        "El ajuste de threshold permitió seleccionar un umbral de 0.55, priorizando el F1-score como métrica de equilibrio. Esta decisión mantiene una precision alta y conserva una capacidad razonable de detección de incumplimientos.",
    )

    add_paragraph_before(document, reference, "Matriz de confusión e interpretabilidad", "Heading 3")
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "confusion_matrix.png",
        "Figura 14: Matriz de confusión final usando Random Forest y threshold optimizado.",
        width_inches=4.8,
    )
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "shap_summary.png",
        "Figura 15: Resumen SHAP global del modelo final.",
    )
    add_picture_before(
        document,
        reference,
        ASSETS_DIR / "shap_waterfall.png",
        "Figura 16: Explicación SHAP local para una predicción individual.",
        width_inches=5.8,
    )
    add_paragraph_before(
        document,
        reference,
        "Las figuras SHAP mantienen la línea interpretativa del proyecto original y permiten observar qué variables aportan más a la predicción. Esto refuerza la explicabilidad del modelo, aspecto especialmente importante en aplicaciones financieras.",
    )

    document.save(DOCX_PATH)
    print(f"Documento actualizado: {DOCX_PATH}")
    print(f"Backup creado: {BACKUP_PATH}")


if __name__ == "__main__":
    main()
