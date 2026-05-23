from __future__ import annotations

from pathlib import Path
import unicodedata

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Documentos" / "acif104_s6_MMedel_2da_Entrega.docx"


def simplify(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(char for char in normalized if not unicodedata.combining(char)).lower()


def find_paragraph(document: Document, needle: str):
    target = simplify(needle)
    for paragraph in document.paragraphs:
        if target in simplify(paragraph.text):
            return paragraph
    raise ValueError(f"No se encontro el parrafo: {needle}")


def replace_text(paragraph, new_text: str) -> None:
    paragraph.clear()
    paragraph.add_run(new_text)


def insert_after(paragraph, text: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.style = paragraph.style
    new_paragraph.add_run(text)
    return new_paragraph


def main() -> None:
    document = Document(DOCX_PATH)

    frontend_paragraph = find_paragraph(
        document,
        "La aplicacion se implemento en Dash",
    )
    replace_text(
        frontend_paragraph,
        "La aplicacion se implemento como un frontend en Dash orientado a usabilidad, con formulario de entrada, validacion de campos, visualizacion de probabilidad, clasificacion de riesgo y secciones especificas para explicabilidad y monitoreo. La interfaz ya no ejecuta el modelo de forma acoplada dentro del mismo flujo, sino que consume servicios HTTP del backend para obtener predicciones, metricas y estado del sistema.",
    )

    backend_paragraph = find_paragraph(
        document,
        "El backend corresponde al pipeline serializado",
    )
    replace_text(
        backend_paragraph,
        "El backend se implemento como una API independiente en Flask. Esta API encapsula el pipeline serializado en app/model.pkl, el threshold optimizado y la logica de inferencia. El flujo actual queda definido como: usuario en Dash -> solicitud HTTP -> API Flask -> pipeline de preprocesamiento y Random Forest -> respuesta JSON -> visualizacion en la interfaz.",
    )

    limitations_paragraph = find_paragraph(
        document,
        "La solucion cumple el alcance demostrativo del proyecto",
    )
    replace_text(
        limitations_paragraph,
        "Con esta refactorizacion, la solucion deja de ser una interfaz monolitica y pasa a una arquitectura desacoplada para el alcance academico del proyecto. El frontend queda separado del backend, la inferencia se expone mediante endpoints dedicados y el sistema incorpora monitoreo persistente de predicciones. La principal limitacion pendiente ya no es la ausencia de backend o monitoreo, sino la falta de mecanismos avanzados como autenticacion, despliegue distribuido y deteccion automatica de drift.",
    )

    inserted = insert_after(
        limitations_paragraph,
        "Los endpoints implementados son GET /health, GET /metrics, POST /predict, GET /monitoring/summary y GET /monitoring/recent. Esta separacion permite auditar el estado del servicio, reutilizar la capa de inferencia y demostrar de forma explicita la distincion entre frontend y backend exigida por la rubrica.",
    )
    insert_after(
        inserted,
        "Para cumplir el requisito de monitoreabilidad, cada prediccion se almacena en SQLite con fecha, variables relevantes, probabilidad y clase estimada. Sobre ese almacenamiento, el frontend consulta volumen de uso, promedio de riesgo, conteo de casos de alto riesgo y ultimas predicciones, materializando un monitoreo basico pero funcional y persistente.",
    )

    improvements_paragraph = find_paragraph(
        document,
        "El modelo presenta un desempeño solido",
    )
    replace_text(
        improvements_paragraph,
        "El modelo presenta un desempeno solido, pero aun mantiene desafios propios del problema: desbalance de clases, presencia de outliers, dependencia del dataset utilizado y ausencia de validacion temporal. A nivel de sistema, la mejora mas relevante ya fue implementada mediante API desacoplada y monitoreo persistente; las mejoras futuras se concentran en autenticacion, despliegue productivo, alertas de drift y trazabilidad historica mas avanzada.",
    )

    repository_paragraph = find_paragraph(
        document,
        "Para reforzar la reproducibilidad, el repositorio incorpora instrucciones",
    )
    replace_text(
        repository_paragraph,
        "Para reforzar la reproducibilidad, el repositorio incorpora instrucciones de instalacion y ejecucion, un script de refinamiento para recalcular metricas y graficos, el pipeline final serializado, el frontend Dash, la API Flask y una base SQLite generada automaticamente para monitoreo. Esto permite revisar el codigo fuente, levantar la arquitectura local completa y verificar tanto la inferencia como el registro persistente de eventos.",
    )

    repository_paragraph_2 = find_paragraph(
        document,
        "El repositorio incluye README con instalacion y ejecucion",
    )
    replace_text(
        repository_paragraph_2,
        "El README actualizado documenta la arquitectura local, los endpoints del backend, la forma de ejecutar el frontend y el comportamiento del monitoreo persistente, lo que facilita la evaluacion tecnica y la reproduccion del sistema.",
    )

    conclusion_paragraph = find_paragraph(
        document,
        "Como cierre de esta segunda iteracion",
    )
    replace_text(
        conclusion_paragraph,
        "Como cierre de esta segunda iteracion, el proyecto evidencia una mejora metodologica y de arquitectura concreta: se amplio la comparacion de modelos, se incorporaron tres arquitecturas MLP, se evaluaron tecnicas de balanceo, se ajusto el threshold final y se desacoplo la solucion en frontend Dash, backend Flask y monitoreo persistente en SQLite. Random Forest se mantiene como modelo final con F1-score 0.8257 y ROC-AUC 0.9292, mientras que la nueva implementacion satisface de mejor forma los requisitos de frontend, backend, explicabilidad y monitoreo solicitados en la evaluacion.",
    )

    followup_paragraph = find_paragraph(
        document,
        "Los resultados obtenidos evidencian que es posible identificar patrones relevantes",
    )
    replace_text(
        followup_paragraph,
        "Los resultados obtenidos evidencian que es posible identificar patrones relevantes en los datos y, al mismo tiempo, encapsular esa capacidad predictiva en una solucion mas cercana a un sistema real: una interfaz de usuario desacoplada, un servicio de inferencia reutilizable y un registro persistente de predicciones para seguimiento operacional.",
    )

    explain_paragraph = find_paragraph(
        document,
        "Asimismo, el uso de tecnicas de interpretabilidad",
    )
    replace_text(
        explain_paragraph,
        "Asimismo, el uso de tecnicas de interpretabilidad como SHAP permite comprender el funcionamiento del modelo, mientras que la separacion por capas y el monitoreo persistente fortalecen la auditabilidad general del sistema. En conjunto, el proyecto no solo cumple los objetivos analiticos planteados, sino que tambien presenta una implementacion mas completa y defendible frente a criterios de evaluacion aplicados a soluciones de machine learning.",
    )

    document.save(DOCX_PATH)
    print(f"Informe actualizado: {DOCX_PATH}")


if __name__ == "__main__":
    main()
