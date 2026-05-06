from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Documentos" / "acif104_s6_MMedel.docx"
OUTPUT_DOCX = ROOT / "Documentos" / "acif104_s6_MMedel_version_excelente.docx"
ASSETS_DIR = ROOT / ".generated_assets"


def fmt(value) -> str:
    try:
        return f"{float(value):.4f}"
    except (TypeError, ValueError):
        return str(value)


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(f"No se encontró el párrafo: {text}")


def find_paragraph_index(document: Document, text: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if text in paragraph.text:
            return index
    raise ValueError(f"No se encontró el párrafo: {text}")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_before(reference_paragraph, element) -> None:
    reference_paragraph._p.addprevious(element)


def add_paragraph_before(document: Document, reference_paragraph, text: str = "", style: str | None = None):
    paragraph = document.add_paragraph(text)
    if style:
        paragraph.style = style
    element = paragraph._p
    element.getparent().remove(element)
    insert_before(reference_paragraph, element)
    return paragraph


def add_picture_before(document: Document, reference_paragraph, image_path: Path, caption: str, width_inches: float = 6.1):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    element = paragraph._p
    element.getparent().remove(element)
    insert_before(reference_paragraph, element)

    caption_paragraph = document.add_paragraph(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_paragraph.runs:
        caption_paragraph.runs[0].italic = True
    caption_element = caption_paragraph._p
    caption_element.getparent().remove(caption_element)
    insert_before(reference_paragraph, caption_element)


def add_table_before(document: Document, reference_paragraph, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)

    element = table._tbl
    element.getparent().remove(element)
    insert_before(reference_paragraph, element)


def remove_section_by_paragraph_text(document: Document, start_text: str, end_text: str) -> None:
    start = find_paragraph_index(document, start_text)
    end = find_paragraph_index(document, end_text)
    if start >= end:
        return
    body = document.element.body
    start_element = document.paragraphs[start]._element
    end_element = document.paragraphs[end]._element
    children = list(body)
    start_xml_index = children.index(start_element)
    end_xml_index = children.index(end_element)
    for element in children[start_xml_index:end_xml_index]:
        body.remove(element)


def replace_paragraph_text(document: Document, contains: str, new_text: str) -> None:
    paragraph = find_paragraph(document, contains)
    paragraph.clear()
    paragraph.add_run(new_text)


def main() -> None:
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)

    comparison = pd.read_csv(ASSETS_DIR / "comparacion_modelos_ml_dl.csv")
    balance = pd.read_csv(ASSETS_DIR / "comparacion_balanceo.csv")
    threshold = pd.read_csv(ASSETS_DIR / "threshold_tuning.csv")
    final_metrics = pd.read_csv(ASSETS_DIR / "metricas_modelo_final.csv").iloc[0]

    document = Document(OUTPUT_DOCX)

    # La evidencia con figuras había quedado después de GitHub. Se elimina de esa ubicación
    # para insertarla donde corresponde: sección V. Resultados.
    if any("Evidencia experimental generada desde el notebook" in p.text for p in document.paragraphs):
        remove_section_by_paragraph_text(document, "Evidencia experimental generada desde el notebook", "CONCLUSIÓN GENERAL")

    replace_paragraph_text(
        document,
        "Además de Regresión Logística, Árbol de Decisión y Random Forest",
        "Además de Regresión Logística, Árbol de Decisión y Random Forest, se incorporó Gradient Boosting como técnica ensemble adicional. También se evaluaron tres configuraciones MLP con scikit-learn: MLP simple con una capa oculta de 32 neuronas, MLP profunda con dos capas ocultas de 64 y 32 neuronas, y MLP regularizada con la misma estructura profunda más penalización L2. En las tres redes se utilizó activación ReLU en capas ocultas, salida probabilística para clasificación binaria y detención temprana con validación interna. La incorporación de estas arquitecturas permite cumplir la comparación ML vs DL sin cambiar el enfoque tabular del proyecto.",
    )

    replace_paragraph_text(
        document,
        "En esta segunda iteración, los modelos MLP sí fueron incorporados",
        "En esta segunda iteración, los modelos MLP fueron incorporados como aproximación neuronal de deep learning mediante scikit-learn. No se seleccionan como arquitectura final porque, para este dataset tabular, Random Forest mantiene mejor equilibrio entre F1-score, ROC-AUC, tiempo de entrenamiento e interpretabilidad. Las redes MLP son técnicamente pertinentes como contraste, pero requieren mayor ajuste y entregan menor transparencia para justificar decisiones de riesgo crediticio.",
    )

    architecture_anchor = find_paragraph(document, "En consecuencia, Random Forest se posiciona como la mejor alternativa")
    add_paragraph_before(document, architecture_anchor, "Refinamiento posterior a la selección inicial", "Heading 3")
    add_paragraph_before(
        document,
        architecture_anchor,
        "Después de la selección inicial se refinó la evaluación mediante comparación ampliada de modelos, técnicas de balanceo y threshold tuning. Para las arquitecturas MLP se consideró la convergencia mediante early stopping, evitando extender entrenamientos sin mejora relevante. En términos de error y métricas, Random Forest mantuvo el mejor F1-score entre los modelos comparados, mientras que el threshold 0.55 permitió ajustar la decisión final sin modificar el entrenamiento del modelo.",
    )
    add_table_before(
        document,
        architecture_anchor,
        ["Arquitectura", "Capas / estructura", "Activación", "Entrada", "Salida", "Rol en el proyecto"],
        [
            ["MLP simple", "1 capa oculta: 32 neuronas", "ReLU", "Variables preprocesadas", "Probabilidad de default", "Baseline neuronal"],
            ["MLP profunda", "2 capas ocultas: 64 y 32 neuronas", "ReLU", "Variables preprocesadas", "Probabilidad de default", "Mayor capacidad no lineal"],
            ["MLP regularizada", "2 capas ocultas: 64 y 32 neuronas + L2", "ReLU", "Variables preprocesadas", "Probabilidad de default", "Control de sobreajuste"],
            ["Random Forest final", "100 árboles de decisión", "No aplica", "Variables crudas vía pipeline", "Probabilidad de default", "Modelo seleccionado"],
        ],
    )

    elaboration_anchor = find_paragraph(document, "En primer lugar, se llevó a cabo la división del dataset")
    add_paragraph_before(
        document,
        elaboration_anchor,
        "La partición principal se realizó en entrenamiento y prueba, manteniendo estratificación sobre la variable objetivo para conservar la proporción de clases. Adicionalmente, las arquitecturas MLP utilizaron validación interna mediante early stopping, por lo que el proceso considera entrenamiento, validación y prueba sin introducir fuga de información. El conjunto de prueba se reservó para la evaluación final de métricas y matriz de confusión.",
    )

    frontend_anchor = find_paragraph(document, "La implementación sigue siendo una demo académica")
    add_paragraph_before(
        document,
        frontend_anchor,
        "Desde el punto de vista de usabilidad, la interfaz permite ingresar variables crudas del solicitante y visualizar una probabilidad comprensible. Desde el punto de vista de escalabilidad, el uso de un pipeline serializado facilita reemplazar o actualizar el modelo sin rehacer la interfaz. El monitoreo se aborda a nivel demostrativo mediante métricas reales generadas por el notebook y leídas desde archivos CSV.",
    )

    results_reference = find_paragraph(document, "PROPUESTA DE MEJORAS")
    add_paragraph_before(document, results_reference, "Resultados experimentales finales", "Heading 3")
    add_paragraph_before(
        document,
        results_reference,
        "Los siguientes resultados fueron generados al ejecutar el notebook completo. Se incorporan aquí para concentrar la evidencia cuantitativa y visual dentro de la sección de resultados, evitando separar las figuras del análisis que las interpreta.",
    )

    add_paragraph_before(document, results_reference, "Métricas finales del modelo seleccionado", "Heading 3")
    add_table_before(
        document,
        results_reference,
        ["Modelo", "Threshold", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [[
            final_metrics["Modelo"],
            fmt(final_metrics["Threshold"]),
            fmt(final_metrics["Accuracy"]),
            fmt(final_metrics["Precision"]),
            fmt(final_metrics["Recall"]),
            fmt(final_metrics["F1"]),
            fmt(final_metrics["ROC_AUC"]),
        ]],
    )

    add_paragraph_before(document, results_reference, "Comparación ML y DL", "Heading 3")
    add_table_before(
        document,
        results_reference,
        ["Modelo", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "Tiempo (s)"],
        [
            [
                row["Modelo"],
                fmt(row["Accuracy"]),
                fmt(row["Precision"]),
                fmt(row["Recall"]),
                fmt(row["F1"]),
                fmt(row["ROC_AUC"]),
                fmt(row["Tiempo_entrenamiento_seg"]),
            ]
            for _, row in comparison.iterrows()
        ],
    )
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "grafico_modelos.png",
        "Figura 11: Comparación de modelos ML y arquitecturas MLP según F1-score.",
    )
    add_paragraph_before(
        document,
        results_reference,
        "Random Forest obtiene el F1-score más alto entre las alternativas evaluadas. Gradient Boosting queda como el modelo más competitivo dentro de los ensembles adicionales. Las arquitecturas MLP muestran desempeño razonable, pero no justifican reemplazar Random Forest debido a su menor interpretabilidad y a que no superan sus métricas principales.",
    )

    add_paragraph_before(document, results_reference, "Balanceo y threshold tuning", "Heading 3")
    add_table_before(
        document,
        results_reference,
        ["Técnica", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            [
                row["Tecnica_balanceo"],
                fmt(row["Precision"]),
                fmt(row["Recall"]),
                fmt(row["F1"]),
                fmt(row["ROC_AUC"]),
            ]
            for _, row in balance.iterrows()
        ],
    )
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "grafico_balanceo.png",
        "Figura 12: Impacto de técnicas de balanceo sobre Random Forest.",
    )
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "threshold_plot.png",
        "Figura 13: Evaluación de thresholds para precision, recall y F1-score.",
    )
    best_threshold = threshold.sort_values(["F1", "Recall"], ascending=False).iloc[0]
    add_paragraph_before(
        document,
        results_reference,
        f"El threshold seleccionado fue {best_threshold['Threshold']:.2f}, ya que maximiza el F1-score en la grilla evaluada. Esto mejora ligeramente el equilibrio general frente al umbral estándar de 0.50, manteniendo alta precision y un recall coherente con el alcance del proyecto.",
    )

    add_paragraph_before(document, results_reference, "Matriz de confusión e interpretabilidad SHAP", "Heading 3")
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "confusion_matrix.png",
        "Figura 14: Matriz de confusión final con Random Forest y threshold optimizado.",
        width_inches=4.8,
    )
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "shap_summary.png",
        "Figura 15: Análisis SHAP global del modelo final.",
    )
    add_picture_before(
        document,
        results_reference,
        ASSETS_DIR / "shap_waterfall.png",
        "Figura 16: Explicación SHAP local para una predicción individual.",
        width_inches=5.8,
    )
    add_paragraph_before(
        document,
        results_reference,
        "El análisis SHAP refuerza el requisito de explicabilidad, ya que permite observar la contribución de variables como carga del préstamo sobre el ingreso, tasa de interés e ingreso del solicitante. Esto conecta las métricas con una lectura interpretable del riesgo crediticio.",
    )

    github_anchor = find_paragraph(document, "https://github.com")
    add_paragraph_before(
        document,
        github_anchor,
        "El repositorio fue reforzado con un README de instalación y ejecución, un script reproducible para recalcular métricas y gráficos, el pipeline final serializado y una carpeta generada automáticamente para artefactos experimentales. Esto mejora la trazabilidad del proyecto y facilita su revisión técnica.",
    )

    conclusion_anchor = find_paragraph(document, "Como cierre de esta segunda iteración")
    replace_paragraph_text(
        document,
        "Como cierre de esta segunda iteración",
        f"Como cierre de esta segunda iteración, el proyecto evidencia una mejora metodológica concreta: se amplió la comparación de modelos, se incorporaron tres arquitecturas MLP, se evaluaron técnicas de balanceo, se ajustó el threshold final y se reforzó la reproducibilidad del pipeline. Random Forest se mantiene como modelo final con F1-score {final_metrics['F1']:.4f} y ROC-AUC {final_metrics['ROC_AUC']:.4f}, no por continuidad automática, sino porque conserva el mejor equilibrio entre desempeño, interpretabilidad y uso práctico sobre datos tabulares.",
    )

    # Ajuste formal básico: fuente uniforme en el contenido agregado y existente.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = None

    document.save(OUTPUT_DOCX)
    print(f"Versión optimizada creada: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
