from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Documentos" / "acif104_s6_MMedel_version_excelente.docx"


def find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(text)


def set_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def insert_after(paragraph, text, style=None):
    new_paragraph = paragraph.insert_paragraph_before("")
    paragraph._p.addnext(new_paragraph._p)
    set_text(new_paragraph, text)
    if style:
        new_paragraph.style = style
    return new_paragraph


doc = Document(DOCX_PATH)

set_text(
    find_paragraph(doc, "En base al análisis anterior, se seleccionaron tres técnicas"),
    "En base al análisis anterior y al refinamiento de la segunda iteración, se seleccionaron modelos de distinta complejidad para su implementación y comparación: tres modelos clásicos de machine learning, un ensemble adicional y tres arquitecturas MLP como aproximación neuronal de deep learning.",
)

rf_paragraph = find_paragraph(doc, "Este modelo es especialmente adecuado para datasets tabulares")
gb_title = insert_after(rf_paragraph, "Gradient Boosting")
gb_title.style = "Normal"
gb_body = insert_after(
    gb_title,
    "Gradient Boosting se incorpora como modelo ensemble adicional para contrastar Random Forest con una técnica basada en árboles secuenciales. Su inclusión permite evaluar si un ensamble de boosting mejora el rendimiento sin perder completamente la compatibilidad con datos tabulares.",
)
mlp_title = insert_after(gb_body, "Arquitecturas MLP")
mlp_title.style = "Normal"
insert_after(
    mlp_title,
    "Se consideran tres configuraciones MLP: una red simple de una capa oculta, una red profunda de dos capas ocultas y una red profunda regularizada. Estas arquitecturas permiten cubrir el componente neuronal solicitado por la rúbrica y comparar su desempeño con modelos de machine learning más interpretables.",
)

set_text(
    find_paragraph(doc, "La elección de estos modelos responde a la necesidad"),
    "La elección de estos modelos responde a la necesidad de comparar distintos enfoques de aprendizaje, desde modelos lineales e interpretables hasta ensambles no lineales y redes neuronales multicapa:",
)
set_text(
    find_paragraph(doc, "Random Forest: modelo avanzado, orientado a maximizar el desempeño"),
    "Random Forest: modelo ensemble seleccionado como arquitectura final por su equilibrio entre desempeño e interpretabilidad",
)
insert_after(
    find_paragraph(doc, "Random Forest: modelo ensemble seleccionado"),
    "Gradient Boosting: modelo ensemble adicional para contrastar desempeño con boosting",
)
insert_after(
    find_paragraph(doc, "Gradient Boosting: modelo ensemble adicional"),
    "MLP simple, MLP profunda y MLP regularizada: arquitecturas neuronales para comparar el enfoque DL solicitado",
)

doc.save(DOCX_PATH)
print(f"Corrección aplicada: {DOCX_PATH}")
