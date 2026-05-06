from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Documentos" / "acif104_s6_MMedel_version_excelente.docx"
OUTPUT = ROOT / "Documentos" / "acif104_s6_MMedel_version_excelente_compacta.docx"


def find_paragraph(document: Document, text: str):
    for paragraph in document.paragraphs:
        if text in paragraph.text:
            return paragraph
    raise ValueError(text)


def body_index(document: Document, paragraph) -> int:
    return list(document.element.body).index(paragraph._element)


def insert_before(reference, text: str, style: str | None = None):
    paragraph = reference.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def remove_between(document: Document, start_text: str, end_text: str, keep_start: bool = True, keep_end: bool = True):
    start = find_paragraph(document, start_text)
    end = find_paragraph(document, end_text)
    body = document.element.body
    children = list(body)
    start_i = body_index(document, start)
    end_i = body_index(document, end)
    remove_start = start_i + 1 if keep_start else start_i
    remove_end = end_i if keep_end else end_i + 1
    for element in children[remove_start:remove_end]:
        body.remove(element)


def replace_text(paragraph, text: str):
    paragraph.clear()
    paragraph.add_run(text)


def main():
    shutil.copy2(SOURCE, OUTPUT)
    document = Document(OUTPUT)

    # Metodología: reemplazo por una versión más directa, conservando el plan de trabajo.
    remove_between(document, "METODOLOGÍA", "Plan de trabajo", keep_start=True, keep_end=True)
    plan = find_paragraph(document, "Plan de trabajo")
    insert_before(
        plan,
        "2.0 Refinamiento iterativo y mejora del proyecto",
        "Heading 2",
    )
    insert_before(
        plan,
        "Esta segunda versión mantiene el enfoque original del proyecto, pero amplía la experimentación para responder a la evaluación previa: incorpora Gradient Boosting, tres arquitecturas MLP, comparación de balanceo, threshold tuning, pipeline reproducible y métricas exportadas desde el notebook.",
    )
    insert_before(
        plan,
        "La metodología se organiza como una adaptación de CRISP-DM: comprensión del problema financiero, análisis exploratorio de datos, selección de técnicas, diseño del pipeline, implementación, validación y comunicación de resultados. El flujo completo evita fuga de información mediante partición estratificada y uso de transformaciones ajustadas solo con datos de entrenamiento.",
    )
    insert_before(
        plan,
        "Las fases principales fueron: EDA para revisar completitud, outliers y patrones; selección de modelos ML y MLP según pertinencia técnica; diseño de arquitectura con preprocesamiento, entrada/salida y métricas; implementación en notebook y Dash; validación con Accuracy, Precision, Recall, F1-score, ROC-AUC, matriz de confusión y SHAP.",
    )

    # Frontend/backend: queda más compacto y evita repetir la idea de demo varias veces.
    remove_between(document, "DESARROLLO DE FRONT END Y BACK END", "RESULTADOS", keep_start=True, keep_end=True)
    results = find_paragraph(document, "RESULTADOS")
    insert_before(
        results,
        "La aplicación se implementó en Dash como una interfaz local para ingresar variables crudas del solicitante y obtener una probabilidad de incumplimiento. El frontend valida campos básicos y muestra la predicción, variables relevantes y métricas de desempeño.",
    )
    insert_before(
        results,
        "El backend corresponde al pipeline serializado en app/model.pkl, que integra imputación, escalamiento, One-Hot Encoding y Random Forest. El flujo de inferencia es: usuario → validación → pipeline → modelo → probabilidad → visualización. Las métricas se leen desde los CSV generados por el notebook, lo que mantiene coherencia entre entrenamiento, evaluación y app.",
    )
    insert_before(
        results,
        "La solución cumple el alcance demostrativo del proyecto, aunque no corresponde a una arquitectura productiva desacoplada: no incluye API independiente, monitoreo persistente ni drift detection. Estas limitaciones se declaran para mantener una evaluación realista del sistema.",
    )

    # Resultados: eliminar bloque antiguo y dejar la evidencia generada por notebook.
    remove_between(document, "RESULTADOS", "Resultados experimentales finales", keep_start=True, keep_end=True)
    results_final = find_paragraph(document, "Resultados experimentales finales")
    insert_before(
        results_final,
        "La evaluación final se realizó con datos no vistos durante el entrenamiento. Se utilizaron métricas de clasificación adecuadas para clases desbalanceadas y se complementó el análisis con matriz de confusión, comparación de modelos, balanceo, threshold tuning y SHAP.",
    )

    # Propuesta de mejoras: compacta, pero mantiene limitaciones y varias mejoras concretas.
    remove_between(document, "PROPUESTA DE MEJORAS", "CÓDIGO FUENTE EN GITHUB", keep_start=True, keep_end=True)
    github = find_paragraph(document, "CÓDIGO FUENTE EN GITHUB")
    insert_before(
        github,
        "El modelo presenta un desempeño sólido, pero mantiene limitaciones propias del problema: desbalance de clases, presencia de outliers, dependencia del dataset utilizado y ausencia de validación temporal o monitoreo en producción. Además, aunque SHAP mejora la explicabilidad, Random Forest sigue siendo menos transparente que un modelo lineal.",
    )
    insert_before(
        github,
        "Una primera mejora consiste en profundizar el tratamiento del desbalance mediante combinaciones de class_weight, undersampling, SMOTE y ajuste de threshold. Esto permitiría controlar con más precisión el trade-off entre precision y recall según el costo de falsos negativos en riesgo crediticio.",
    )
    insert_before(
        github,
        "Una segunda mejora es optimizar hiperparámetros con validación cruzada, Random Search o Grid Search, evaluando número de árboles, profundidad máxima y variables consideradas por división. Esto podría mejorar la generalización sin cambiar el enfoque principal del proyecto.",
    )
    insert_before(
        github,
        "También se propone ampliar el feature engineering y comparar modelos tabulares adicionales como XGBoost o LightGBM. Su incorporación debe evaluarse manteniendo el equilibrio entre desempeño, interpretabilidad, costo computacional y facilidad de documentación.",
    )

    # Cierre GitHub más breve.
    try:
        replace_text(
            find_paragraph(document, "El repositorio fue reforzado con un README"),
            "El repositorio incluye README con instalación y ejecución, notebook, app Dash, script reproducible, pipeline serializado y artefactos generados para respaldar la revisión técnica.",
        )
    except ValueError:
        pass

    document.save(OUTPUT)
    print(f"Versión compacta creada: {OUTPUT}")


if __name__ == "__main__":
    main()
