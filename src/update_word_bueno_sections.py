from __future__ import annotations

import unicodedata
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Documentos" / "acif104_s6_MMedel_2da_Entrega.docx"


def simplify(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    return "".join(
        character for character in normalized if not unicodedata.combining(character)
    ).lower()


def find_paragraph(document: Document, needle: str):
    target = simplify(needle)
    for paragraph in document.paragraphs:
        if target in simplify(paragraph.text):
            return paragraph
    raise ValueError(f"No se encontro el parrafo: {needle}")


def replace_text(paragraph, new_text: str) -> None:
    paragraph.clear()
    paragraph.add_run(new_text)


def main() -> None:
    document = Document(DOCX_PATH)

    replace_text(
        find_paragraph(document, "Usabilidad: El sistema debe ser fácil de utilizar"),
        "Usabilidad: El sistema debe ser fácil de utilizar, permitiendo que usuarios no técnicos puedan ingresar datos y comprender los resultados sin dificultad. En esta versión, la interfaz fue reorganizada en una vista más clara, con tarjetas, mejor separación visual, controles numéricos más cómodos, una zona principal para el resultado y apoyos visuales como explicación de variables, gráfico de desempeño e historial reciente.",
    )

    replace_text(
        find_paragraph(document, "Monitoreabilidad: El sistema debe permitir evaluar continuamente su desempeño"),
        "Monitoreabilidad: El sistema debe permitir seguir el comportamiento general de la solución a medida que se utiliza. Para responder a este criterio, se incorporó un registro persistente de predicciones y un panel que resume cantidad de consultas, promedio de riesgo, casos clasificados como alto riesgo e historial reciente.",
    )

    replace_text(
        find_paragraph(document, "La aplicacion se implemento como un frontend en Dash"),
        "El frontend se desarrolló en Dash como una aplicación de una sola vista, pensada para que el uso sea directo y fácil de entender. El usuario puede ingresar los datos del solicitante, ajustar valores numéricos de forma más cómoda, visualizar la probabilidad estimada de incumplimiento y recibir una clasificación clara entre bajo y alto riesgo.",
    )

    replace_text(
        find_paragraph(document, "El backend se implemento como una API independiente en Flask"),
        "En paralelo, el backend se separó en una API construida con Flask. Esta capa recibe la información enviada desde la interfaz, aplica el pipeline entrenado y devuelve la predicción al frontend, evitando que toda la lógica quede mezclada en la misma aplicación visual.",
    )

    replace_text(
        find_paragraph(document, "Con esta refactorizacion, la solucion deja de ser una interfaz monolitica"),
        "Esta separación mejora el orden general de la solución y permite distinguir con mayor claridad la parte visual de la parte encargada de procesar la predicción. En términos prácticos, esto hace que la aplicación resulte más consistente, más fácil de mantener y más defendible frente al criterio de frontend y backend solicitado en la rúbrica.",
    )

    replace_text(
        find_paragraph(document, "Los endpoints implementados son GET /health"),
        "Además de la predicción individual, la interfaz incorpora una sección visual de apoyo con las métricas más relevantes del modelo. En lugar de limitarse a mostrar texto, esta vista presenta un gráfico de desempeño y una explicación breve de las variables con mayor influencia, lo que vuelve la interacción más clara y más útil para quien revisa el sistema.",
    )

    replace_text(
        find_paragraph(document, "Para cumplir el requisito de monitoreabilidad"),
        "Junto con lo anterior, se agregó un monitoreo básico pero funcional. Cada predicción queda almacenada y luego se resume en el panel de la aplicación, permitiendo observar cuántas consultas se han realizado, cuál es el riesgo promedio estimado y cuáles fueron las predicciones más recientes. Con ello, la solución deja de ser solo una demo visual y pasa a mostrar también seguimiento del uso.",
    )

    replace_text(
        find_paragraph(document, "Para reforzar la reproducibilidad, el repositorio incorpora instrucciones"),
        "Para facilitar la revisión, el repositorio reúne el notebook, el modelo entrenado, el frontend en Dash, la API en Flask y los archivos necesarios para ejecutar la solución de forma local. De esta manera, quien evalúe el proyecto no solo puede leer el informe, sino también revisar cómo está organizada la implementación y reproducir su funcionamiento.",
    )

    replace_text(
        find_paragraph(document, "El README actualizado documenta la arquitectura local"),
        "Además, el README fue actualizado con instrucciones de instalación y ejecución más claras, de modo que la puesta en marcha del sistema no dependa de interpretaciones adicionales. Esto fortalece la calidad de las referencias técnicas del proyecto, porque deja mejor documentado cómo acceder al código y cómo probar la aplicación.",
    )

    replace_text(
        find_paragraph(document, "El repositorio incluye los notebooks desarrollados"),
        "En conjunto, estas referencias permiten seguir el trabajo desde el análisis en notebook hasta la aplicación final, facilitando la trazabilidad del proyecto y mejorando la presentación formal del código fuente entregado.",
    )

    document.save(DOCX_PATH)
    print(f"Documento actualizado: {DOCX_PATH}")


if __name__ == "__main__":
    main()
